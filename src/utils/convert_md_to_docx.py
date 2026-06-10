import os
import re
import subprocess
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_paragraph_borders(p):
    """Adds a subtle border to a paragraph (for code blocks)."""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_page_number(run):
    """Adds a page number field to a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def convert_docx_to_pdf(docx_path, pdf_path):
    """Uses macOS AppleScript to compile DOCX to PDF using Microsoft Word."""
    abs_docx_path = os.path.abspath(docx_path)
    abs_pdf_path = os.path.abspath(pdf_path)
    
    print(f"Compiling PDF from {abs_docx_path} using Microsoft Word...")
    
    applescript = f'''
    tell application "Microsoft Word"
        set docPath to POSIX file "{abs_docx_path}"
        set pdfPath to POSIX file "{abs_pdf_path}"
        open docPath
        set activeDoc to active document
        save as activeDoc file name pdfPath file format format PDF
        close activeDoc saving no
    end tell
    '''
    
    try:
        process = subprocess.run(
            ['osascript', '-e', applescript],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=45
        )
        if process.returncode == 0:
            print(f"PDF successfully generated at: {abs_pdf_path}")
        else:
            print(f"PDF conversion failed: {process.stderr}")
    except Exception as e:
        print(f"Error during PDF conversion: {e}")

def parse_markdown_to_docx(md_path, docx_path):
    print(f"Reading markdown from: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Define standard margins (2.54 cm or 1 inch standard medium margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Enable different headers and footers for the first page (cover page)
        section.different_first_page_header_footer = True
        
        # Configure standard page header
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Rapport de Stage BUT2 | Milan LOI | RGU Social Media Analysis"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Calibri'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(120, 120, 120)

        # Configure standard page footer (with page numbering)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), 2)  # Right tab stop at 6.5 inches
        
        run_text = fp.add_run("Milan LOI | ÉTUDE DE LA POLARISATION DISCURSIVE\t")
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(8.5)
        run_text.font.color.rgb = RGBColor(120, 120, 120)
        
        run_page = fp.add_run()
        run_page.font.name = 'Calibri'
        run_page.font.size = Pt(8.5)
        run_page.font.color.rgb = RGBColor(120, 120, 120)
        add_page_number(run_page)

    # Configure base font style to Calibri 11
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    in_code_block = False
    code_content = []
    
    in_table = False
    table_rows = []
    
    cover_page = True

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Skip Table of Contents placeholder in markdown (automated TOC will be injected)
        if "L'IUT de Limoges exige" in stripped or (stripped.startswith("[") and stripped.endswith("]")):
            idx += 1
            continue

        # Handle code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                # Shading equivalent
                run = p.add_run("\n".join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(50, 50, 50)
                add_paragraph_borders(p)
                
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            idx += 1
            continue

        if in_code_block:
            code_content.append(line.rstrip('\n'))
            idx += 1
            continue

        # Handle Markdown Tables
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            idx += 1
            continue
        elif in_table:
            # Table has ended
            # Parse table_rows
            header_row = [cell.strip() for cell in table_rows[0].split('|')[1:-1]]
            # Skip alignment row (contains ---)
            data_rows = []
            for row in table_rows[2:]:
                cells = [cell.strip() for cell in row.split('|')[1:-1]]
                if cells:
                    data_rows.append(cells)
            
            # Create DOCX table
            if header_row:
                t = doc.add_table(rows=len(data_rows) + 1, cols=len(header_row))
                t.style = 'Light Shading Accent 1'
                
                # Populate Header
                hdr_cells = t.rows[0].cells
                for col_idx, text in enumerate(header_row):
                    hdr_cells[col_idx].text = text
                    hdr_cells[col_idx].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)
                    set_cell_background(hdr_cells[col_idx], '3B82F6') # Sapphire Blue
                    hdr_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                
                # Populate Data
                for r_idx, row_cells in enumerate(data_rows):
                    row_obj = t.rows[r_idx + 1]
                    for c_idx, val in enumerate(row_cells):
                        if c_idx < len(row_obj.cells):
                            clean_val = val.replace('**', '').replace('*', '')
                            row_obj.cells[c_idx].text = clean_val
                            row_obj.cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
                            if r_idx % 2 == 1:
                                set_cell_background(row_obj.cells[c_idx], 'F3F4F6')
                
                # Add trailing space after table
                doc.add_paragraph().paragraph_format.space_after = Pt(10)
            
            table_rows = []
            in_table = False
            continue

        # Handle page breaks (only separate the cover page, ignore thematic breaks in body)
        if stripped == "---" or stripped == '***':
            if cover_page:
                # Toggle cover page flag, but do NOT add a page break here.
                # The first heading (# REMERCIEMENTS) will automatically add a page break,
                # preventing a blank page.
                cover_page = False
            idx += 1
            continue

        # Ignore manual HTML page breaks to let layout flow naturally
        if stripped == '<div style="page-break-after: always;"></div>':
            idx += 1
            continue

        # Skip empty lines
        if not stripped:
            idx += 1
            continue

        # Handle Headings
        if stripped.startswith("#"):
            match = re.match(r'^(#+)\s+(.*)$', stripped)
            if match:
                level = len(match.group(1))
                title = match.group(2)
                
                # Heading 1 starts on a new page/section (except main title on the cover page)
                if level == 1 and not cover_page:
                    # Specific requirement: pagination stops at the conclusion.
                    # Section 6 (self-assessment) and subsequent chapters are in a new section without page numbers.
                    if title.startswith("6. AUTO-ÉVALUATION"):
                        new_section = doc.add_section()
                        new_section.footer.is_linked_to_previous = False
                        fp = new_section.footer.paragraphs[0]
                        fp.text = ""
                        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run_text = fp.add_run("Milan LOI | ÉTUDE DE LA POLARISATION DISCURSIVE")
                        run_text.font.name = 'Calibri'
                        run_text.font.size = Pt(8.5)
                        run_text.font.color.rgb = RGBColor(120, 120, 120)
                        
                        new_section.header.is_linked_to_previous = True
                    else:
                        doc.add_page_break()
                
                if level == 1:
                    heading_size = Pt(18)
                    heading_bold = True
                elif level == 2:
                    heading_size = Pt(14)
                    heading_bold = True
                else:
                    heading_size = Pt(12)
                    heading_bold = True

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                
                run = p.add_run(title)
                run.bold = heading_bold
                run.font.size = heading_size
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(59, 130, 246) # Sapphire Blue
                
                if level == 1 and cover_page:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.size = Pt(22)
                
                # If we just added the SOMMAIRE heading, inject the Word field code for an automated TOC
                if title == "SOMMAIRE":
                    toc_p = doc.add_paragraph()
                    toc_p.paragraph_format.space_after = Pt(12)
                    run_toc = toc_p.add_run()
                    
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'separate')
                    fldChar3 = OxmlElement('w:fldChar')
                    fldChar3.set(qn('w:fldCharType'), 'end')
                    
                    r = run_toc._r
                    r.append(fldChar1)
                    r.append(instrText)
                    r.append(fldChar2)
                    r.append(fldChar3)

            idx += 1
            continue

        # Handle Paragraphs and list items
        is_list = stripped.startswith("* ") or stripped.startswith("- ") or (stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or stripped.startswith("4. ") or stripped.startswith("5. "))
        
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0 # Single line spacing
        p.paragraph_format.space_after = Pt(6)
        
        if is_list:
            p.paragraph_format.left_indent = Inches(0.4)
            if stripped.startswith("* ") or stripped.startswith("- "):
                bullet_run = p.add_run("•  ")
                bullet_run.bold = True
                content = stripped[2:]
            else:
                num_match = re.match(r'^(\d+\.)\s+(.*)$', stripped)
                if num_match:
                    num_run = p.add_run(f"{num_match.group(1)}  ")
                    num_run.bold = True
                    content = num_match.group(2)
                else:
                    content = stripped
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            content = stripped
            
            # Format centered equations
            if content.startswith("$$") and content.endswith("$$"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                content = content.replace("$$", "")
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                
            if cover_page:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

        # Parse inline markdown formatting (bold **, italic *)
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', content)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                r = p.add_run(token[2:-2])
                r.bold = True
            elif token.startswith("*") and token.endswith("*"):
                r = p.add_run(token[1:-1])
                r.italic = True
            elif token.startswith("$") and token.endswith("$"):
                r = p.add_run(token[1:-1])
                r.font.name = 'Consolas'
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(120, 20, 20)
            else:
                p.add_run(token)

        idx += 1

    # Save Document
    print(f"Saving compiled DOCX report to: {docx_path}")
    doc.save(docx_path)
    print("DOCX successfully generated.")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        md_path = sys.argv[1]
        docx_path = sys.argv[2]
    else:
        md_path = "docs/Rapport_de_Stage_Milan_Loi.md"
        docx_path = "docs/Rapport_de_Stage_Milan_Loi.docx"
        
    parse_markdown_to_docx(md_path, docx_path)
    
    # Compile to PDF using Word AppleScript
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert_docx_to_pdf(docx_path, pdf_path)
